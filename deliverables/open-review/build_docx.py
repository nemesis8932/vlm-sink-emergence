#!/usr/bin/env python3
"""sections-unslop/*.md -> paper-v2-ste-unslop.docx (text content; figures become captions).

Usage: python3 build_docx.py            # unslop sections
       python3 build_docx.py --stock    # the original sections/ set
Word cannot place the SVG figures, so each <figure> renders as a bracketed placeholder
plus its caption. Tables, headings, lists and inline emphasis carry over.
"""
import re, sys, html
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
STOCK = "--stock" in sys.argv
SRC = HERE / ("sections" if STOCK else "sections-unslop")
OUT = HERE / ("paper-v2-ste.docx" if STOCK else "paper-v2-ste-unslop.docx")
TITLE = "Attention-Sink Signatures Dissociate During Vision–Language Pretraining"
SUBTITLE = "Anonymous author(s) · Submitted to VLM4RWD @ NeurIPS 2026"

INLINE = re.compile(r"(\*\*.+?\*\*|(?<![\w*])\*[^*\n]+?\*(?![\w*])|`[^`]+?`)")
SUPSUB = re.compile(r"Sink\^([0-9.]+|ε)(?:_1)?")


def add_runs(par, text):
    """**bold**, *italic*, `code`, plus <sup>/<sub> and Sink^x_1."""
    text = text.replace("<sup>", "\x01").replace("</sup>", "\x02")
    text = text.replace("<sub>", "\x03").replace("</sub>", "\x04")
    text = SUPSUB.sub(lambda m: f"Sink\x01{m.group(1)}\x02\x031\x04", text)
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        bold = italic = mono = False
        if chunk.startswith("**") and chunk.endswith("**"):
            chunk, bold = chunk[2:-2], True
        elif chunk.startswith("*") and chunk.endswith("*"):
            chunk, italic = chunk[1:-1], True
        elif chunk.startswith("`") and chunk.endswith("`"):
            chunk, mono = chunk[1:-1], True
        # split out sup/sub markers inside the chunk
        for piece in re.split(r"(\x01[^\x02]*\x02|\x03[^\x04]*\x04)", chunk):
            if not piece:
                continue
            sup = piece.startswith("\x01")
            sub = piece.startswith("\x03")
            body = piece.strip("\x01\x02\x03\x04")
            r = par.add_run(html.unescape(body))
            r.bold, r.italic = bold, italic
            if mono:
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            if sup:
                r.font.superscript = True
            if sub:
                r.font.subscript = True


def flush(doc, buf, style=None, indent=None):
    if not buf:
        return
    par = doc.add_paragraph(style=style)
    if indent is not None:
        par.paragraph_format.left_indent = Inches(indent)
    add_runs(par, " ".join(buf))
    buf.clear()


def add_table(doc, rows):
    grid = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    head, body = grid[0], grid[2:]
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Table Grid"
    for cell, txt in zip(t.rows[0].cells, head):
        cell.paragraphs[0].clear()
        add_runs(cell.paragraphs[0], txt)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in body:
        cells = t.add_row().cells
        for cell, txt in zip(cells, row):
            cell.paragraphs[0].clear()
            add_runs(cell.paragraphs[0], txt)
    doc.add_paragraph()


def figure_block(doc, raw):
    alt = re.search(r'alt="([^"]*)"', raw)
    cap = re.search(r"<figcaption>(.*?)</figcaption>", raw, re.S)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(f"[figure omitted: {alt.group(1) if alt else 'figure'}]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    if cap:
        text = re.sub(r"<[^>]+>", "", cap.group(1))
        text = " ".join(text.split())
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, text)
        for run in p.runs:
            run.font.size = Pt(9)


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(10.5)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(10)

    for path in sorted(SRC.glob("[0-9][0-9]-*.md")):
        lines = path.read_text().split("\n")
        buf, i = [], 0
        while i < len(lines):
            ln = lines[i]
            if ln.startswith("<figure"):
                flush(doc, buf)
                block = []
                while i < len(lines) and "</figure>" not in lines[i]:
                    block.append(lines[i]); i += 1
                block.append(lines[i] if i < len(lines) else "")
                figure_block(doc, "\n".join(block))
                i += 1
                continue
            if ln.lstrip().startswith("|") and i + 1 < len(lines) and set(lines[i+1].replace("|", "").strip()) <= set("-: "):
                flush(doc, buf)
                rows = []
                while i < len(lines) and lines[i].lstrip().startswith("|"):
                    rows.append(lines[i]); i += 1
                add_table(doc, rows)
                continue
            if re.match(r"^\s*</?(div|p|br|span|table|hr)\b[^>]*>\s*$", ln):
                flush(doc, buf)   # bare HTML block wrappers (e.g. the references div)
                i += 1
                continue
            m = re.match(r"^(#{1,3}) (.*)", ln)
            if m:
                flush(doc, buf)
                doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
                i += 1
                continue
            if re.match(r"^\s*(\d+\.|-)\s+", ln):
                flush(doc, buf)
                item = [re.sub(r"^\s*(\d+\.|-)\s+", "", ln)]
                i += 1
                while i < len(lines) and lines[i].startswith("   ") and lines[i].strip():
                    item.append(lines[i].strip()); i += 1
                flush(doc, item, indent=0.25)
                continue
            if not ln.strip():
                flush(doc, buf)
                i += 1
                continue
            buf.append(ln.strip())
            i += 1
        flush(doc, buf)

    doc.save(OUT)
    print(f"wrote {OUT.name}")


main()
